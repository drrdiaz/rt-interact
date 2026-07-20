from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/rodrigodiaz/Documents/(11) Projects/(02) Prototype/(06) Cowork Prototype/rt-interact")
RAW_JSON = ROOT / "src/data/raw/app-output-resolved.json"
OUT_DIR = ROOT / "audit-output"
OUT_PATH = OUT_DIR / "RT_Interact_Database_Repair_Audit_2026-07-18.docx"

REPAIRED_SITES = [
    "brain",
    "lung",
    "oesophagus",
    "pelvis",
    "head_neck",
    "spine",
    "upper_abdomen",
    "breast",
    "skin",
    "bone",
    "mediastinum",
    "tbi",
]
LEGACY_SITES = ["nonspecific"]

SITE_NOTES = {
    "brain": "CNS-scoped toxicity language restored; unsupported direct labels downgraded; fallback messaging no longer implies safety when exact citations are missing.",
    "lung": "Thoracic outputs foreground pulmonary/oesophageal risk and stereotactic caveats for central/ultracentral versus peripheral treatments.",
    "oesophagus": "Rows foreground dysphagia, mucosal injury, and fistula/vascular concern instead of generic class-toxicity phrasing.",
    "pelvis": "Rows foreground bowel, GU, marrow, and fistula/perforation issues rather than undifferentiated systemic toxicity bundles.",
    "head_neck": "Rows foreground mucosal, swallowing, salivary, airway, skin, and carotid toxicity; generic no-data outputs were replaced with anatomy-scoped warnings.",
    "spine": "Rows foreground cord/cauda, vertebral, and marrow toxicity instead of cross-site class toxicity bundles.",
    "upper_abdomen": "Rows foreground hepatic/RILD, upper-GI, pancreatic/biliary, renal, and marrow toxicity instead of vague abdominal wording.",
    "breast": "Rows foreground skin/chest wall, pulmonary, cardiac, breast/soft-tissue, and lymphoedema toxicity; immune and HER2 rows now read coherently.",
    "skin": "Rows foreground dermatitis, wound-healing, recall, and soft-tissue breakdown rather than mixed-site class outputs.",
    "bone": "Rows foreground marrow, fracture, and soft-tissue risk, with osteonecrosis/healing flags for bone-modifying agents.",
    "mediastinum": "Legacy thoracic rows now read as pulmonary, oesophageal, cardiac, marrow, and endocrine mediastinal risk rather than generic site-generic toxicity.",
    "tbi": "Legacy total-body-irradiation rows now read as multi-organ marrow/pulmonary/hepatic/renal/ocular/gonadal risk rather than generic no-data outputs for checked agents.",
}

EXAMPLE_SCENARIOS = {
    ("brain", "cisplatin", "Palliative", "Concurrent"): "Cisplatin + brain + palliative RT",
    ("lung", "abemaciclib", "Stereotactic", "Concurrent or sequential"): "Abemaciclib + thoracic RT + stereotactic RT",
    ("oesophagus", "cisplatin", "Conventional", "Concurrent"): "Cisplatin + oesophageal RT + conventional RT",
    ("pelvis", "cisplatin", "Conventional", "Concurrent"): "Cisplatin + pelvic RT + conventional RT",
    ("head_neck", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + head and neck RT + conventional RT",
    ("spine", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + spine RT + conventional RT",
    ("upper_abdomen", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + upper-abdominal RT + conventional RT",
    ("breast", "abemaciclib", "Conventional", "Concurrent or sequential"): "Abemaciclib + breast/chest-wall RT + conventional RT",
    ("skin", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + superficial RT + conventional RT",
    ("bone", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + bone RT + conventional RT",
    ("mediastinum", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + mediastinal RT + conventional RT",
    ("tbi", "pembrolizumab", "Conventional", "Concurrent or sequential"): "Pembrolizumab + TBI + conventional RT",
}


def load_rows() -> list[dict]:
    return json.loads(RAW_JSON.read_text())


def compute_site_stats(rows: list[dict]) -> dict[str, dict[str, int]]:
    by_site: dict[str, dict[str, int]] = {}
    for row in rows:
        site = row["site_key"]
        by_site.setdefault(
            site,
            {
                "rows": 0,
                "directish": 0,
                "direct_empty": 0,
                "no_applicable": 0,
                "generic_drug_level": 0,
                "evidence_not_loaded": 0,
            },
        )
        stats = by_site[site]
        stats["rows"] += 1
        if row["evidence_directness"] in {"Direct RT-combination", "Partially direct RT-combination"}:
            stats["directish"] += 1
            if not str(row.get("citations") or "").strip():
                stats["direct_empty"] += 1
        if row["evidence_directness"] == "No applicable evidence":
            stats["no_applicable"] += 1
        if row["evidence_directness"] == "Drug-level (site-generic)":
            stats["generic_drug_level"] += 1
        if row["evidence_directness"] == "Evidence source not loaded":
            stats["evidence_not_loaded"] += 1
    return by_site


def get_example_rows(rows: list[dict]) -> list[dict]:
    examples = []
    for (site, agent, frac, timing), label in EXAMPLE_SCENARIOS.items():
        match = next(
            (
                row
                for row in rows
                if row["site_key"] == site
                and agent in row["searchable_agents"]
                and row["fractionation_key"] == frac
                and row["timing_key"] == timing
            ),
            None,
        )
        if match:
            examples.append(
                {
                    "scenario": label,
                    "site": site,
                    "risk": match["risk_category"],
                    "warning": match["interaction_warning"],
                    "toxicity": match["toxicity_domains"],
                    "directness": match["evidence_directness"],
                }
            )
    return examples


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 11, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("RT-Interact Database Repair Audit")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "Repair summary for the deployed static feed after site-taxonomy simplification and block-by-block scenario-scope correction."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10.5)

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date(2026, 7, 18).isoformat()}").bold = True


def add_summary(doc: Document, site_stats: dict[str, dict[str, int]]) -> None:
    doc.add_heading("Executive Summary", level=1)
    bullets = [
        "The active simplified clinician-facing site blocks have now been repaired across brain, lung, oesophagus, pelvis, head and neck, spine, upper abdomen, breast, skin, and bone.",
        "Legacy non-UI blocks were also cleaned for mediastinum and total-body irradiation so the static feed reads coherently beyond the simplified front-end taxonomy.",
        "Empty-citation direct-claim rows were cleared from all repaired active site blocks and from repaired legacy thoracic/TBI spot checks.",
        "The main remaining residual issue is legacy `nonspecific` timing-only material plus any future row-by-row evidence refinement, not broad taxonomy failure.",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Current Feed Snapshot", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "Site key",
        "Rows",
        "Direct/partial direct",
        "Direct with empty citations",
        "No applicable evidence",
        "Drug-level generic",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    for site in REPAIRED_SITES:
        stats = site_stats[site]
        row = table.add_row().cells
        values = [
            site,
            str(stats["rows"]),
            str(stats["directish"]),
            str(stats["direct_empty"]),
            str(stats["no_applicable"]),
            str(stats["generic_drug_level"]),
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value)


def add_repaired_blocks(doc: Document, site_stats: dict[str, dict[str, int]], examples: list[dict]) -> None:
    doc.add_heading("Repaired Blocks", level=1)
    for site in REPAIRED_SITES:
        doc.add_heading(site.replace("_", " ").title(), level=2)
        doc.add_paragraph(SITE_NOTES[site])
        stats = site_stats[site]
        doc.add_paragraph(
            f"Rows: {stats['rows']}. Direct/partial-direct rows remaining: {stats['directish']}. "
            f"Direct rows with empty citations: {stats['direct_empty']}. "
            f"No-applicable-evidence rows: {stats['no_applicable']}. "
            f"Drug-level generic rows: {stats['generic_drug_level']}."
        )

        site_examples = [example for example in examples if example["site"] == site]
        if site_examples:
            table = doc.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ["Scenario", "Alert", "Main warning", "Toxicities", "Evidence label"]
            for cell, header in zip(table.rows[0].cells, headers):
                set_cell_text(cell, header, bold=True)
            for example in site_examples:
                row = table.add_row().cells
                values = [
                    example["scenario"],
                    example["risk"],
                    example["warning"],
                    example["toxicity"],
                    example["directness"],
                ]
                for cell, value in zip(row, values):
                    set_cell_text(cell, value)


def add_remaining_work(doc: Document, site_stats: dict[str, dict[str, int]]) -> None:
    doc.add_heading("Remaining Legacy Material", level=1)
    doc.add_paragraph(
        "The main remaining residual material is now the legacy `nonspecific` site bucket. It is timing-oriented fallback logic rather than an active clinician-facing site block."
    )

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Site key", "Rows", "Direct with empty citations", "No applicable evidence", "Drug-level generic", "Priority"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    for site in LEGACY_SITES:
        stats = site_stats[site]
        row = table.add_row().cells
        values = [
            site,
            str(stats["rows"]),
            str(stats["direct_empty"]),
            str(stats["no_applicable"]),
            str(stats["generic_drug_level"]),
            "Later",
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value)

    doc.add_heading("Key Residual Note", level=2)
    doc.add_paragraph(
        "The `nonspecific` block still contains agent/timing heuristics such as recall, washout, and class-level stereotactic escalation. It does not drive the simplified site-based front end in the same way as the grouped site blocks."
    )


def add_next_steps(doc: Document) -> None:
    doc.add_heading("Recommended Next Steps", level=1)
    steps = [
        "Decide whether legacy `nonspecific` rows should be retained as hidden fallback logic or retired from the deployed static feed.",
        "If the prototype is now clinically acceptable, freeze a reviewed database snapshot and regenerate the static feed from that reviewed state rather than continuing ad hoc row edits.",
        "Regenerate any downstream documentation or clinician-facing validation sheets from the repaired feed so review artefacts match the deployed prototype.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def build_document() -> Path:
    rows = load_rows()
    site_stats = compute_site_stats(rows)
    examples = get_example_rows(rows)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_summary(doc, site_stats)
    doc.add_page_break()
    add_repaired_blocks(doc, site_stats, examples)
    doc.add_page_break()
    add_remaining_work(doc, site_stats)
    add_next_steps(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
